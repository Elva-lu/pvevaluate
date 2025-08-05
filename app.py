from flask import Flask, request, send_file
import os
import io
import re
from docx import Document

app = Flask(__name__)

TEMPLATE_PATH = "CC2025016_MA評估回饋20250417.docx"

# 提前載入 Word 模板邏輯文字，避免每次請求都重新開檔
def load_logic_text():
    if not os.path.exists(TEMPLATE_PATH):
        return None, "⚠️ 找不到範本檔案"
    doc = Document(TEMPLATE_PATH)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text, f"依 {os.path.basename(TEMPLATE_PATH)} 邏輯"

LOGIC_TEXT, STATUS_MESSAGE = load_logic_text()


# 區塊抽取
def extract_between(text, start_marker, end_marker):
    pattern = re.escape(start_marker) + r'(.*?)' + re.escape(end_marker)
    match = re.search(pattern, text, re.DOTALL)
    return match.group(1).strip() if match else ""

# 解析 AI 產出的長文本
def parse_ai_report(text):
    return {
        "part_number": re.search(r"Case ID: (\S+)", text).group(1) if re.search(r"Case ID: (\S+)", text) else "無資料",
        "adverse_event": extract_between(text, "**2. 不良反應事件:**", "**3. 可疑藥物:**"),
        "suspected_drug": extract_between(text, "**3. 可疑藥物:**", "**4. 併用藥物:**"),
        "case_summary": extract_between(text, "**個案摘要:**", "**1. 病人資料:**"),
        "patient_info": extract_between(text, "**1. 病人資料:**", "**2. 不良反應事件:**"),
        "concomitant_medications": extract_between(text, "**4. 併用藥物:**", "**5. 藥物相關性評估:**"),
        "drug_relationship_assessment": extract_between(text, "**5. 藥物相關性評估:**", "**6. 文獻支持:**"),
        "literature_support": extract_between(text, "**6. 文獻支持:**", "**7. 結論:**"),
        "conclusion": extract_between(text, "**7. 結論:**", "**8. 建議:**"),
        "recommendation": extract_between(text, "**8. 建議:**", "**備註:**"),
        "notes": extract_between(text, "**備註:**", ""),
    }

# 建立 Word 檔案並填入欄位
def generate_word_report(data, logic_text, status_message):
    doc = Document()
    doc.add_heading("藥物不良反應評估報告", level=1)
    doc.add_paragraph(f"📌 {status_message}")

    # 開頭三項
    doc.add_paragraph(f"🔢 料品號: {data['part_number']}")
    doc.add_paragraph("📍 不良反應事件:")
    doc.add_paragraph(data["adverse_event"])
    doc.add_paragraph("💊 可疑藥物:")
    doc.add_paragraph(data["suspected_drug"])

    # 背景
    doc.add_heading("🧾 背景", level=2)
    doc.add_paragraph(data["case_summary"])

    # 評估回饋
    doc.add_heading("📋 評估回饋", level=2)
    combined = "\n\n".join([
        data["patient_info"],
        data["concomitant_medications"],
        data["drug_relationship_assessment"],
        data["conclusion"],
        data["recommendation"]
    ])
    doc.add_paragraph(combined.strip())

    # 備註
    doc.add_heading("📝 備註", level=2)
    doc.add_paragraph(data["notes"])

    # 參考資料
    doc.add_heading("📚 參考資料", level=2)
    doc.add_paragraph(data["literature_support"])

    # 存成 BytesIO
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


@app.route("/", methods=["GET"])
def home():
    return "✅ Flask 服務運作中，請 POST 到 /evaluate 上傳報告文字"

@app.route("/evaluate", methods=["POST"])
def evaluate():
    raw_text = request.form.get("text")
    if not raw_text:
        return "❌ 缺少欄位 text，請提供 AI 產出的完整報告", 400

    if not LOGIC_TEXT:
        return STATUS_MESSAGE, 500

    try:
        # 解析 AI 內容
        parsed = parse_ai_report(raw_text)
        # 產出 Word 檔案
        report = generate_word_report(parsed, LOGIC_TEXT, STATUS_MESSAGE)
        filename = f"ADR_Report_{parsed['part_number']}.docx"

        response = send_file(report, as_attachment=True, download_name=filename)
        response.headers["X-Status-Message"] = STATUS_MESSAGE
        return response

    except Exception as e:
        return f"❌ 產生報告失敗: {str(e)}", 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, threaded=True)
